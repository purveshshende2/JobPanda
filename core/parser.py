"""Resume/JD file parsing: PDF, DOCX, TXT, MD."""
import io
import os


def parse_file(data: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext in (".txt", ".md", ""):
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, TXT or MD.")


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf is required to read PDFs. pip install pypdf") from e

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError(
            "Could not extract text from this PDF. It may be a scanned image - "
            "export a text-based PDF, DOCX, or paste your resume directly."
        )
    return text


def _parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError("python-docx is required to read DOCX files.") from e

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(l for l in lines if l.strip())
