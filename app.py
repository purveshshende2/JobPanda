"""JobPanda - AI Job Application Assistant (Streamlit UI)."""
import io
import os
import re

from dotenv import find_dotenv, load_dotenv

load_dotenv()
load_dotenv(find_dotenv(usecwd=False))

import streamlit as st

from core.analyzer import analyze, analyze_resume_only
from core.llm import LLMClient
from core.parser import parse_file
from core.render import resume_to_latex, resume_to_pdf
from core.tailor import cover_letter, deep_review, tailor_resume

ICON = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "panda.png")
st.set_page_config(page_title="JobPanda", page_icon=ICON, layout="wide")


def md_to_docx(md_text: str) -> bytes:
    """Very small markdown -> docx converter (headings + bullets + paragraphs)."""
    import docx
    from docx.shared import Pt

    d = docx.Document()
    style = d.styles["Normal"]
    style.font.size = Pt(10.5)
    for line in md_text.splitlines():
        line = line.rstrip()
        if re.match(r"^#\s", line):
            h = d.add_heading(line[2:].strip(), level=1)
            h.alignment = 1
        elif re.match(r"^#{2,6}\s", line):
            d.add_heading(re.sub(r"^#+\s*", "", line), level=2)
        elif re.match(r"^\s*[-*•]\s+", line):
            p = d.add_paragraph(re.sub(r"^\s*[-*•]\s+", "", line), style="List Bullet")
            _clean_runs(p)
        elif not line.strip():
            continue
        else:
            p = d.add_paragraph(line)
            _clean_runs(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _clean_runs(paragraph):
    for run in paragraph.runs:
        run.text = run.text.replace("*", "").replace("#", "")

# ---------- state ----------
for key, default in {
    "resume_text": "",
    "jd_text": "",
    "analysis": None,
    "tailored": None,
    "letter": None,
    "quick_resume_text": "",
    "quick_result": None,
    "deep_review": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ---------- sidebar ----------
with st.sidebar:
    st.title("🐼 JobPanda")
    st.caption("Resume → ATS score → tailored resume → apply")
    st.divider()

    st.subheader("AI settings (optional)")
    PRESETS = {
        "Groq — Qwen3.6-27B (free tier)": (
            "https://api.groq.com/openai/v1", "qwen/qwen3.6-27b"),
        "Groq — GPT-OSS 120B (free tier)": (
            "https://api.groq.com/openai/v1", "openai/gpt-oss-120b"),
        "OpenAI": ("https://api.openai.com/v1", "gpt-4o-mini"),
        "NVIDIA NIM — Nemotron (free tier)": (
            "https://integrate.api.nvidia.com/v1", "nvidia/nemotron-3-super-120b-a12b"),
        "OpenRouter — Nemotron :free": (
            "https://openrouter.ai/api/v1", "nvidia/nemotron-3-super-120b-a12b:free"),
        "Ollama (local)": ("http://localhost:11434/v1", "llama3.1"),
    }

    def apply_preset():
        b, m = PRESETS[st.session_state.preset]
        st.session_state["base_url_box"] = b
        st.session_state["model_box"] = m

    st.selectbox("Provider preset", list(PRESETS), key="preset", on_change=apply_preset,
                 help="Picking a preset auto-fills Base URL & Model below.")
    api_key = st.text_input("API key", value=os.getenv("OPENAI_API_KEY", ""), type="password",
                            help="Stays in this browser session's memory only - never written to disk, "
                                 "never sent anywhere except the provider you chose.")
    base_url = st.text_input("Base URL", value=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
                             key="base_url_box")
    model = st.text_input("Model", value=os.getenv("OPENAI_MODEL", "gpt-4o-mini"), key="model_box")
    st.caption(f"Raw model ID — e.g. `{PRESETS[st.session_state.preset][1]}`")

    if st.button("🔌 Test AI connection", use_container_width=True):
        c = LLMClient.from_env(api_key=api_key or None, base_url=base_url or None, model=model or None)
        if not c.configured:
            st.warning("Enter an API key first.")
        elif "—" in c.model:
            st.error(f"'{c.model}' is a preset label, not a model ID. "
                     f"Paste this instead: `{PRESETS[st.session_state.preset][1]}`")
        else:
            # provider-mismatch guard
            base_l = c.base_url.lower()
            hints = []
            if c.api_key.startswith("gsk_") and "groq" not in base_l:
                hints.append("Key is a **Groq** key (gsk_...) but Base URL is not Groq. Use `https://api.groq.com/openai/v1`")
            if c.api_key.startswith("nvapi-") and "nvidia" not in base_l:
                hints.append("Key is an **NVIDIA** key (nvapi-...). Use `https://integrate.api.nvidia.com/v1`")
            if c.api_key.startswith("sk-or-") and "openrouter" not in base_l:
                hints.append("Key is an **OpenRouter** key. Use `https://openrouter.ai/api/v1`")
            if "/" in c.model and "openai.com" in base_l:
                hints.append("Model IDs with '/' don't exist on OpenAI — wrong provider for this model.")
            if hints:
                st.error("⚠️ **Provider mismatch:**\n\n- " + "\n\n- ".join(hints))
            else:
                with st.spinner("Pinging provider..."):
                    try:
                        reply = c.chat("You are a connectivity test.", "Reply with exactly: OK", max_tokens=10)
                        st.success(f"Connected · {c.model} · replied: {reply[:40]}")
                    except Exception as e:
                        msg = str(e)
                        st.error(f"Connection failed: {msg[:400]}")
                        # model retired/renamed? offer one-click fix from live catalog
                        if "model_not_found" in msg or "does not exist" in msg.lower():
                            try:
                                ids = c.list_models()
                                chat_ids = [i for i in ids if not any(
                                    t in i.lower() for t in ("whisper", "guard", "tts", "embed"))]
                                st.info(f"Your key works! Click a model below to use it:")
                                for mid in chat_ids[:12]:
                                    if st.button(f"📋 {mid}", key=f"pick_{mid}"):
                                        st.session_state["model_box"] = mid
                                        st.rerun()
                            except Exception as e2:
                                st.caption(f"(Couldn't list models: {str(e2)[:150]})")

    def get_client():
        c = LLMClient.from_env(api_key=api_key or None, base_url=base_url or None, model=model or None)
        return c

    st.caption("Needed only for Tailoring & Cover Letter. ATS scoring is offline.")


def llm_or_warn():
    client = get_client()
    if not client.configured:
        st.warning("Add your API key in the sidebar to use AI features. Get one at platform.openai.com.")
        return None
    return client


# ---------- tabs ----------
tab_upload, tab_score, tab_tailor, tab_letter, tab_quick = st.tabs(
    ["📄 1. Resume & JD", "🔍 2. ATS Score", "✍️ 3. Tailored Resume",
     "📝 4. Cover Letter", "⚡ Quick ATS Check (no JD)"]
)

# ================= TAB 1: upload =================
with tab_upload:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Your resume")
        file = st.file_uploader("Upload PDF / DOCX / TXT (auto-parses)", type=["pdf", "docx", "txt", "md"])
        if file is not None:
            sig = f"{file.name}|{len(file.getvalue())}"
            if st.session_state.get("resume_sig") != sig:
                try:
                    st.session_state.resume_text = parse_file(file.getvalue(), file.name)
                    st.session_state.resume_sig = sig
                except Exception as e:
                    st.error(str(e))
            if st.session_state.get("resume_sig") == sig:
                st.caption(f"✅ Parsed **{file.name}** · {len(st.session_state.resume_text.split())} words")
        if st.session_state.resume_text != st.session_state.get("resume_seeded"):
            st.session_state["resume_box"] = st.session_state.resume_text
            st.session_state.resume_seeded = st.session_state.resume_text
        st.session_state.resume_text = st.text_area(
            "...or paste/edit resume text", value=st.session_state.resume_text,
            height=300, key="resume_box")

    with col2:
        st.subheader("Job description")
        if st.button("Load sample JD"):
            with open("sample_data/sample_job.txt", encoding="utf-8") as f:
                st.session_state.jd_text = f.read()
            st.session_state["jd_box"] = st.session_state.jd_text
            st.session_state.jd_seeded = st.session_state.jd_text
            st.rerun()
        if st.session_state.jd_text != st.session_state.get("jd_seeded"):
            st.session_state["jd_box"] = st.session_state.jd_text
            st.session_state.jd_seeded = st.session_state.jd_text
        st.session_state.jd_text = st.text_area(
            "Paste full job description", value=st.session_state.jd_text,
            height=340, key="jd_box")

    r_ok = len(st.session_state.resume_text.strip()) > 80
    j_ok = len(st.session_state.jd_text.strip()) > 100
    ready = r_ok and j_ok
    hint = "" if ready else ("Add your resume text first." if not r_ok else "Add the job description first.")
    if st.button("🚀 Run ATS Analysis", type="primary", use_container_width=True,
                 disabled=not ready, help=hint):
        with st.spinner("Scoring..."):
            st.session_state.analysis = analyze(st.session_state.resume_text, st.session_state.jd_text)
            st.session_state.tailored = None
            st.session_state.letter = None
            st.session_state.deep_review = None
        st.success(f"✅ Analysis complete — score {st.session_state.analysis['score']}/100."
                   " Now open the **🔍 2. ATS Score** tab above.")
    if not ready:
        st.info(hint)

# ================= TAB 2: score =================
with tab_score:
    a = st.session_state.analysis
    if not a:
        st.info("Run the analysis from tab 1 first.")
    else:
        c1, c2, c3 = st.columns([1, 2, 2])
        with c1:
            st.metric("ATS Score", f"{a['score']}/100", a["grade"])
            if a["score"] < 70:
                st.error(f"Grade: {a['grade']} - keep going!")
            elif a["score"] < 85:
                st.warning(f"Grade: {a['grade']} - close!")
            else:
                st.success(f"Grade: {a['grade']} - apply now!")
        with c2:
            for label, (pct, weight) in a["breakdown"].items():
                st.progress(pct / 100, text=f"{label} — {pct}%")
        with c3:
            st.markdown("**Formatting checks**")
            for check, ok in a["formatting_checks"].items():
                st.write(("✅" if ok else "❌"), check)

        st.divider()
        mcol1, mcol2 = st.columns(2)
        with mcol1:
            st.markdown(f"### ✅ Matched keywords ({len(a['matched_keywords'])})")
            if a["matched_keywords"]:
                st.write(" ".join(f"`{k}`" for k in a["matched_keywords"]))
        with mcol2:
            st.markdown(f"### 🔑 Missing keywords ({len(a['missing_keywords'])})")
            if a["missing_keywords"]:
                st.write(" ".join(f":red[{k}]" for k in a["missing_keywords"]))

        if a.get("years_required"):
            st.markdown(f"**Experience:** JD wants ~{a['years_required']} yrs · resume shows {a['years_in_resume']}")
        if a.get("education_required"):
            st.markdown(f"**Education mentioned in JD:** {', '.join(a['education_required'])}")

        st.divider()
        st.markdown("### 💡 Suggestions to raise your score")
        for tip in a["suggestions"]:
            st.markdown(f"- {tip}")

        # ---- hybrid layer: optional LLM deep review (numeric score stays algorithmic) ----
        st.divider()
        st.markdown("### 🧠 AI Deep Review (optional)")
        st.caption("LLM adds bullet-by-bullet critique & semantic gap analysis on top of the "
                   "algorithmic score above. Needs an API key in the sidebar.")
        if st.button("🧠 Run AI Deep Review"):
            client = llm_or_warn()
            if client:
                with st.spinner("Reviewing every bullet... (~60s)"):
                    try:
                        st.session_state.deep_review = deep_review(
                            client, st.session_state.resume_text, st.session_state.jd_text,
                            a["score"], a["missing_keywords"])
                    except Exception as e:
                        st.error(f"LLM call failed: {e}")
        if st.session_state.deep_review:
            st.markdown(st.session_state.deep_review)
            st.download_button("⬇️ Download review (.md)", st.session_state.deep_review.encode(),
                               "ai_deep_review.md")

# ================= TAB 3: tailor =================
with tab_tailor:
    if not st.session_state.analysis:
        st.info("Run the ATS analysis first so tailoring can target your gaps.")
    elif st.session_state.tailored:
        st.success("Tailored resume generated - review every line for truthfulness before sending!")
        edited = st.text_area("Tailored resume (editable)", value=st.session_state.tailored, height=500)
        try:
            pdf_bytes = resume_to_pdf(edited)
        except Exception as e:
            pdf_bytes = None
            st.warning(f"PDF export failed: {e}")
        try:
            tex_text = resume_to_latex(edited)
        except Exception as e:
            tex_text = ""
            st.warning(f"LaTeX export failed: {e}")
        dl1, dl2, dl3, dl4 = st.columns(4)
        dl1.download_button("⬇️ Download .pdf", pdf_bytes or b"", "tailored_resume.pdf", type="primary",
                            disabled=not pdf_bytes)
        dl2.download_button("⬇️ Download .tex (LaTeX)", tex_text.encode(), "tailored_resume.tex",
                            disabled=not tex_text,
                            help="Upload to overleaf.com and compile with pdflatex for a pixel-perfect PDF.")
        dl3.download_button("⬇️ Download .md", edited.encode(), "tailored_resume.md")
        docx_bytes = md_to_docx(edited)
        dl4.download_button("⬇️ Download .docx", docx_bytes, "tailored_resume.docx")
    else:
        st.markdown("Rewrites your resume against this specific JD using your real experience only.")
        if st.button("✍️ Generate tailored resume", type="primary"):
            client = llm_or_warn()
            if client:
                with st.spinner("Rewriting your resume... (~30s)"):
                    try:
                        st.session_state.tailored = tailor_resume(
                            client, st.session_state.resume_text, st.session_state.jd_text)
                        st.rerun()
                    except Exception as e:
                        st.error(f"LLM call failed: {e}")

# ================= TAB 4: cover letter =================
with tab_letter:
    if not st.session_state.jd_text:
        st.info("Add a job description in tab 1 first.")
    elif st.session_state.letter:
        st.success("Cover letter ready!")
        edited = st.text_area("Cover letter (editable)", value=st.session_state.letter, height=400)
        st.download_button("⬇️ Download .txt", edited.encode(), "cover_letter.txt", type="primary")
    else:
        cc1, cc2 = st.columns(2)
        company = cc1.text_input("Company name (optional)")
        role = cc2.text_input("Role title (optional)")
        if st.button("📝 Write my cover letter", type="primary"):
            client = llm_or_warn()
            if client:
                with st.spinner("Writing..."):
                    try:
                        st.session_state.letter = cover_letter(
                            client, st.session_state.resume_text, st.session_state.jd_text,
                            company=company, role=role)
                        st.rerun()
                    except Exception as e:
                        st.error(f"LLM call failed: {e}")

# ================= TAB 5: quick ATS check (no JD) =================
with tab_quick:
    st.subheader("⚡ General ATS-friendliness check")
    st.caption("Scores formatting & structure only — no job description needed. "
               "For job-specific matching use tab 1 → 2.")

    qfile = st.file_uploader("Upload resume (auto-parses)", type=["pdf", "docx", "txt", "md"], key="quick_file")
    if qfile is not None:
        qsig = f"{qfile.name}|{len(qfile.getvalue())}"
        if st.session_state.get("quick_sig") != qsig:
            try:
                st.session_state.quick_resume_text = parse_file(qfile.getvalue(), qfile.name)
                st.session_state.quick_sig = qsig
                st.session_state.quick_result = None
            except Exception as e:
                st.error(str(e))
        if st.session_state.get("quick_sig") == qsig:
            st.caption(f"✅ Parsed **{qfile.name}** · {len(st.session_state.quick_resume_text.split())} words")
    if st.session_state.quick_resume_text != st.session_state.get("quick_seeded"):
        st.session_state["quick_box"] = st.session_state.quick_resume_text
        st.session_state.quick_seeded = st.session_state.quick_resume_text
    st.session_state.quick_resume_text = st.text_area(
        "...or paste resume text", value=st.session_state.quick_resume_text,
        height=260, key="quick_box")

    qc1, qc2 = st.columns([1, 1])
    with qc1:
        run_quick = st.button("⚡ Score my resume", type="primary",
                              disabled=len(st.session_state.quick_resume_text.strip()) < 80,
                              help="Needs at least ~80 characters of resume text")
    with qc2:
        if st.button("➡️ Use this resume in the full JD analysis"):
            st.session_state.resume_text = st.session_state.quick_resume_text
            st.session_state.resume_seeded = st.session_state.resume_text
            st.success("Copied to tab 1 — add the JD there and run the analysis.")

    if run_quick:
        st.session_state.quick_result = analyze_resume_only(st.session_state.quick_resume_text)

    qr = st.session_state.quick_result
    if qr:
        b1, b2 = st.columns([1, 2])
        with b1:
            st.metric("Resume ATS Score", f"{qr['score']}/100", qr["grade"])
            if qr["score"] < 70:
                st.error(qr["grade"])
            elif qr["score"] < 85:
                st.warning(qr["grade"])
            else:
                st.success(qr["grade"])
        with b2:
            for label, ok, weight, _tip in qr["checks"]:
                st.write(("✅" if ok else "❌"), f"{label} `+{weight}`")

        st.divider()
        st.markdown("### 💡 How to improve")
        for tip in qr["suggestions"]:
            st.markdown(f"- {tip}")
